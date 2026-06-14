"""Word document conversion utilities."""

from pathlib import Path

from docx import Document


def word_to_pdf(input_path: str, output_path: str) -> str:
    """Convert a Word document (.docx) to PDF.

    Uses python-docx for reading and reportlab for PDF generation.
    For complex formatting, LibreOffice CLI is preferred (see word_to_pdf_cli).

    Args:
        input_path: Source .docx file path.
        output_path: Target PDF file path.

    Returns:
        Absolute path of the output PDF.
    """
    # Try LibreOffice first (better fidelity)
    try:
        return word_to_pdf_cli(input_path, output_path)
    except Exception:
        pass

    # Fallback: python-docx + reportlab
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import mm

    src = Path(input_path)
    dst = Path(output_path)
    dst.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(src))
    pdf_doc = SimpleDocTemplate(str(dst), pagesize=A4,
                                leftMargin=25 * mm, rightMargin=25 * mm,
                                topMargin=25 * mm, bottomMargin=25 * mm)

    styles = getSampleStyleSheet()
    story = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            story.append(Spacer(1, 6 * mm))
            continue

        style_name = "Normal"
        if para.style and para.style.name:
            sn = para.style.name.lower()
            if "heading 1" in sn or "title" in sn:
                style_name = "Heading1"
            elif "heading 2" in sn:
                style_name = "Heading2"
            elif "heading 3" in sn:
                style_name = "Heading3"

        story.append(Paragraph(text, styles[style_name]))

    pdf_doc.build(story)
    return str(dst.resolve())


def word_to_pdf_cli(input_path: str, output_path: str) -> str:
    """Convert Word to PDF using LibreOffice CLI (high fidelity).

    Requires: libreoffice installed on the system.
    """
    import subprocess

    src = Path(input_path)
    dst = Path(output_path)
    dst.parent.mkdir(parents=True, exist_ok=True)

    result = subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(dst.parent), str(src)],
        capture_output=True, text=True, timeout=60
    )

    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")

    # LibreOffice saves with same name but .pdf extension
    generated = dst.parent / f"{src.stem}.pdf"
    if generated.exists() and str(generated) != str(dst):
        generated.rename(dst)

    return str(dst.resolve())


def extract_text_from_word(input_path: str) -> str:
    """Extract plain text from a Word document.

    Args:
        input_path: Source .docx file path.

    Returns:
        Extracted text content.
    """
    src = Path(input_path)
    if not src.exists():
        raise FileNotFoundError(f"Input file not found: {src}")

    doc = Document(str(src))
    return "\n".join(para.text for para in doc.paragraphs)
